#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Parse the chemistry docx and output full structure as JSON."""

import sys
import os
import json

# Fix encoding for Windows console
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document

docx_path = r'F:/PBL项目学习/小程序开发/ESAT题库上线/ESAT_Chemistry_65_Questions_题库整理_v3.docx'
doc = Document(docx_path)

print(f"Total paragraphs: {len(doc.paragraphs)}")
print(f"Total tables: {len(doc.tables)}")
print()

results = []

for ti, table in enumerate(doc.tables):
    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(cells)
    results.append(rows)

# Output all tables
for ti, rows in enumerate(results):
    print(f"=== Table {ti} ===")
    for ri, cells in enumerate(rows):
        print(f"  R{ri}: {json.dumps(cells, ensure_ascii=False)}")
    print()
