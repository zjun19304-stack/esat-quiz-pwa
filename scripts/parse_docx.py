#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Parse ESAT biology docx and extract question content."""

from docx import Document
import json

doc_path = r'F:/PBL项目学习/小程序开发/ESAT题库上线/ESAT_Guide_Biology_June2025ESAT生物模拟题库.docx'
doc = Document(doc_path)

# Extract all paragraphs with their styles
lines = []
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text:
        lines.append({
            'idx': i,
            'style': para.style.name,
            'text': text
        })

# Also extract tables (some questions might be in tables)
tables_content = []
for ti, table in enumerate(doc.tables):
    rows_data = []
    for ri, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]
        rows_data.append(cells)
    if rows_data:
        tables_content.append({'table_idx': ti, 'rows': rows_data})

print(f"=== PARAGRAPHS ({len(lines)} non-empty) ===")
for line in lines:
    print(f"[{line['idx']:04d}] ({line['style']}) {line['text'][:200]}")

print(f"\n=== TABLES ({len(tables_content)}) ===")
for t in tables_content:
    print(f"\nTable {t['table_idx']}:")
    for r in t['rows']:
        print(f"  {r}")
